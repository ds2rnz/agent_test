from __future__ import annotations
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
import unicodedata
from typing import Final
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from pypdf import PdfReader


MAX_PDF_SIZE: Final[int] = 50 * 1024 * 1024
DEFAULT_FONT_NAME: Final[str] = "맑은 고딕"

_SAFE_FILENAME_PATTERN: Final[re.Pattern[str]] = re.compile(
    r'[<>:"/\\|?*\x00-\x1f]')

_WINDOWS_RESERVED_NAMES: Final[set[str]] = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


class PdfConversionError(ValueError):
    """사용자에게 안내할 수 있는 PDF 변환 오류."""


@dataclass(frozen=True)
class PdfConversionResult:
    """한 PDF의 변환 결과."""

    source_name: str
    docx_name: str
    txt_name: str
    docx_bytes: bytes
    txt_bytes: bytes
    page_count: int
    text_page_count: int


def _safe_stem(file_name: str) -> str:
    """
    다운로드 파일명에 사용할 안전한 파일명을 만듭니다.

    Windows에서 사용할 수 없는 문자와 예약어를 처리하고,
    지나치게 긴 파일명을 제한합니다.
    """
    stem = Path(file_name).stem.strip() or "converted_pdf"
    stem = unicodedata.normalize("NFC", stem)

    stem = _SAFE_FILENAME_PATTERN.sub("_", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    stem = stem.rstrip(" .")

    if not stem:
        stem = "converted_pdf"

    if stem.upper() in _WINDOWS_RESERVED_NAMES:
        stem = f"_{stem}"

    return stem[:120]


def _normalize_text(text: str) -> str:
    """
    PDF 추출 과정에서 생기는 과도한 공백과 줄바꿈을 정리합니다.

    단, PDF 원문 구조가 훼손되지 않도록 줄 내부의 공백은
    과도하게 줄이지 않습니다.
    """
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = text.replace("\x00", "")

    lines = [line.rstrip() for line in text.splitlines()]
    text = "\n".join(lines)

    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _validate_pdf_bytes(pdf_bytes: bytes) -> None:
    """PDF 바이트 입력값의 기본 유효성을 확인합니다."""
    if not pdf_bytes:
        raise PdfConversionError("PDF 파일이 비어 있습니다.")

    if len(pdf_bytes) > MAX_PDF_SIZE:
        raise PdfConversionError("PDF 파일은 50MB 이하만 변환할 수 있습니다.")


def extract_pdf_pages(pdf_bytes: bytes, password: str | None = None) -> list[str]:
    """
    PDF의 각 페이지에서 텍스트를 추출합니다.

    Args:
        pdf_bytes: PDF 파일 바이트
        password: 암호화 PDF인 경우 사용할 비밀번호.
                  None이면 빈 비밀번호로 해제를 시도합니다.

    Returns:
        페이지별 추출 텍스트 목록

    Raises:
        PdfConversionError: 사용자가 이해할 수 있는 변환 오류
    """
    _validate_pdf_bytes(pdf_bytes)

    try:
        reader = PdfReader(BytesIO(pdf_bytes), strict=False)
    except Exception as error:
        raise PdfConversionError(
            "PDF를 열 수 없습니다. 손상되었거나 지원되지 않는 형식인지 확인해 주세요."
        ) from error

    if reader.is_encrypted:
        try:
            decrypt_result = reader.decrypt(password or "")
        except Exception as error:
            raise PdfConversionError(
                "암호화된 PDF입니다. 올바른 비밀번호를 입력하거나 암호를 해제해 주세요."
            ) from error

        if not decrypt_result:
            raise PdfConversionError(
                "암호화된 PDF입니다. 올바른 비밀번호를 입력하거나 암호를 해제해 주세요."
            )

    try:
        page_count = len(reader.pages)
    except Exception as error:
        raise PdfConversionError(
            "PDF 페이지 정보를 읽을 수 없습니다. 파일이 손상되었는지 확인해 주세요."
        ) from error

    if page_count == 0:
        raise PdfConversionError("PDF에 페이지가 없습니다.")

    pages: list[str] = []

    for page_index in range(page_count):
        try:
            page = reader.pages[page_index]
            text = page.extract_text() or ""
            pages.append(_normalize_text(text))
        except Exception:
            # 일부 페이지에서만 텍스트 추출에 실패하는 경우
            # 전체 변환을 중단하지 않고 빈 페이지로 처리합니다.
            pages.append("")

    if not any(page.strip() for page in pages):
        raise PdfConversionError(
            "PDF에서 텍스트를 찾지 못했습니다. "
            "스캔 이미지 PDF라면 OCR 처리가 먼저 필요합니다."
        )

    return pages


def _build_txt(pages: list[str]) -> bytes:
    """
    페이지별 텍스트를 TXT 파일 바이트로 생성합니다.

    utf-8-sig는 Windows 메모장에서 한글이 깨지는 문제를 줄여 줍니다.
    """
    sections: list[str] = []

    for page_number, text in enumerate(pages, start=1):
        content = text.strip() or "[이 페이지에서 추출된 텍스트가 없습니다.]"
        sections.append(f"===== {page_number}페이지 =====\n{content}")

    return ("\n\n".join(sections) + "\n").encode("utf-8-sig")


def _set_korean_font(style, font_name: str = DEFAULT_FONT_NAME) -> None:
    """
    python-docx 스타일에 한글 폰트를 안정적으로 적용합니다.

    style.font.name만 지정하면 영문 폰트만 적용되고,
    한글에는 적용되지 않는 경우가 있어 eastAsia 폰트를 별도 지정합니다.
    """
    style.font.name = font_name

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts

    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:eastAsia"), font_name)


def _apply_default_document_style(document: Document) -> None:
    """DOCX 문서의 기본 여백과 글꼴 스타일을 설정합니다."""
    section = document.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    for style_name in ("Normal", "Title", "Heading 1"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue

        _set_korean_font(style)

    normal_style = document.styles["Normal"]
    normal_style.font.size = Pt(10.5)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15


def _build_docx(source_name: str, pages: list[str]) -> bytes:
    """페이지별 텍스트를 DOCX 파일 바이트로 생성합니다."""
    document = Document()
    _apply_default_document_style(document)

    document.core_properties.title = Path(source_name).name
    document.core_properties.subject = "PDF 변환 문서"

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.add_run(Path(source_name).name)

    document.add_paragraph(f"PDF 변환 문서 · 총 {len(pages)}페이지")

    for page_number, text in enumerate(pages, start=1):
        if page_number > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        document.add_heading(f"{page_number}페이지", level=1)

        if not text.strip():
            paragraph = document.add_paragraph(
                "[이 페이지에서 추출된 텍스트가 없습니다.]"
            )
            if paragraph.runs:
                paragraph.runs[0].italic = True
            continue

        for line in text.splitlines():
            document.add_paragraph(line if line else " ")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def convert_pdf(
    pdf_bytes: bytes,
    file_name: str,
    password: str | None = None,
) -> PdfConversionResult:
    """
    PDF 하나를 DOCX와 TXT로 변환합니다.

    Args:
        pdf_bytes: PDF 파일 바이트
        file_name: 원본 PDF 파일명
        password: 암호화 PDF인 경우 사용할 비밀번호.
                  기존처럼 비밀번호 없이 사용해도 됩니다.

    Returns:
        PdfConversionResult

    Raises:
        PdfConversionError: 사용자에게 안내 가능한 변환 오류
    """
    if Path(file_name).suffix.lower() != ".pdf":
        raise PdfConversionError("PDF 파일만 변환할 수 있습니다.")

    pages = extract_pdf_pages(pdf_bytes, password=password)
    stem = _safe_stem(file_name)

    return PdfConversionResult(
        source_name=file_name,
        docx_name=f"{stem}.docx",
        txt_name=f"{stem}.txt",
        docx_bytes=_build_docx(file_name, pages),
        txt_bytes=_build_txt(pages),
        page_count=len(pages),
        text_page_count=sum(bool(page.strip()) for page in pages),
    )